import streamlit as st
import google.generativeai as genai
from datetime import datetime, timedelta
import time
from docx import Document
import io
from supabase import create_client, Client
import extra_streamlit_components as stx
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="ChemMind AI - Pro", page_icon="🧪", layout="wide")

# --- ÖZEL STİL (CSS) ---
st.markdown("""
    <style>
    .stAlert { border-radius: 10px; }
    .report-box { 
        background-color: #ffffff; 
        padding: 20px; 
        border: 1px solid #dee2e6; 
        border-radius: 15px;
        border-left: 8px solid #28a745;
    }
    
    /* En alttaki 'Made with Streamlit' yazısını gizle */
    footer {visibility: hidden;}
    
    /* Streamlit Cloud İnatçı Butonlarını Yok Etme (Sadece 3 Nokta Kalacak) */
    .stDeployButton {display: none !important;}
    [data-testid="manage-app-button"] {display: none !important;}
    [data-testid="stToolbarActionButton"] {display: none !important;}
    [data-testid="stToolbar"] a {display: none !important;}
    .viewerBadge_container__1QSob {display: none !important;}
    .viewerBadge_link__1S137 {display: none !important;}
    
    </style>
    """, unsafe_allow_html=True)

# --- ÇEREZ (COOKIE) YÖNETİCİSİ ---
cookie_manager = stx.CookieManager()

# --- VERİTABANI (SUPABASE) BAĞLANTISI ---
@st.cache_resource
def init_supabase() -> Client:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase = init_supabase()

# --- API AYARLARI ---
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except:
    API_KEY = ""
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# --- OTURUM (SESSION) HAFIZALARI ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_info" not in st.session_state:
    st.session_state.user_info = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "current_subject" not in st.session_state:
    st.session_state.current_subject = "Genel"
if "force_logout" not in st.session_state:
    st.session_state.force_logout = False

# --- MAİL GÖNDERME FONKSİYONU ---
def mail_gonder(doc_buffer, ogrenci_ad, konu):
    try:
        sender = st.secrets["EMAIL_SENDER"]
        password = st.secrets["EMAIL_PASSWORD"]
        receiver = st.secrets["HOCA_EMAIL"]
        
        msg = MIMEMultipart()
        msg['From'] = sender
        msg['To'] = receiver
        msg['Subject'] = f"ChemMind AI: {ogrenci_ad} - {konu} Raporu"
        
        govde = f"Merhaba Hocam,\n\n{ogrenci_ad} adlı öğrencinin {konu} konusu üzerine gerçekleştirdiği ChemMind AI etkileşim raporu ektedir.\n\nİyi çalışmalar dileriz."
        msg.attach(MIMEText(govde, 'plain'))
        
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(doc_buffer.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{ogrenci_ad}_Raporu.docx"')
        msg.attach(part)
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender, password)
        server.send_message(msg)
        server.quit()
        return True, "Rapor başarıyla e-postalandı!"
    except Exception as e:
        return False, f"Mail Hatası: {str(e)}"

def sifre_hatirlat_mail(alici_mail, sifre, ad):
    try:
        sender = st.secrets["EMAIL_SENDER"]
        password = st.secrets["EMAIL_PASSWORD"]
        
        msg = MIMEMultipart()
        msg['From'] = sender
        msg['To'] = alici_mail
        msg['Subject'] = "ChemMind AI - Şifre Hatırlatma"
        
        govde = f"Merhaba {ad},\n\nChemMind AI laboratuvarı için mevcut şifreniz: {sifre}\n\nLütfen şifrenizi kimseyle paylaşmayın. İyi çalışmalar dileriz!"
        msg.attach(MIMEText(govde, 'plain'))
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender, password)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        return False

# --- OTOMATİK GİRİŞ KONTROLÜ ---
cerez_ogrenci_no = cookie_manager.get(cookie="chem_user")
if cerez_ogrenci_no and not st.session_state.logged_in and not st.session_state.force_logout:
    res = supabase.table("kullanicilar").select("*").eq("ogrenci_no", cerez_ogrenci_no).execute()
    if len(res.data) > 0:
        st.session_state.logged_in = True
        st.session_state.user_info = res.data[0]
        st.rerun()

# --- GİRİŞ VE KAYIT EKRANI ---
if not st.session_state.logged_in:
    st.title("🧪 ChemMind AI Laboratuvarına Hoş Geldiniz")
    st.markdown("Lütfen laboratuvara girmek için öğrenci kimliğinizi doğrulayın.")
    
    tab1, tab2 = st.tabs(["🔑 Giriş Yap", "📝 Yeni Kayıt Ol"])
    with tab1:
        with st.form("login_form"):
            l_email = st.text_input("E-Posta Adresiniz:") # l_no yerine l_email oldu
            l_pass = st.text_input("Şifreniz:", type="password")
            if st.form_submit_button("Laboratuvara Gir", use_container_width=True):
                # Sorgu artık e-posta üzerinden yapılıyor
                res = supabase.table("kullanicilar").select("*").eq("eposta", l_email.lower()).eq("sifre", l_pass).execute()
                
                if len(res.data) > 0:
                    st.session_state.logged_in = True
                    st.session_state.user_info = res.data[0]
                    st.session_state.force_logout = False 
                    
                    # Otomatik giriş çerezi için yine benzersiz olan ogrenci_no'yu kullanalım
                    cookie_manager.set("chem_user", res.data[0]["ogrenci_no"], expires_at=datetime.now() + timedelta(days=30))
                    
                    st.success("Giriş başarılı! Laboratuvara hoş geldiniz.")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Hatalı e-posta adresi veya şifre!")
    with tab2:
        with st.form("register_form"):
            r_no = st.text_input("Öğrenci Numaranız (Varsa):")
            r_name = st.text_input("Adınız ve Soyadınız:")
            r_email = st.text_input("E-Posta Adresiniz (Giriş için):")
            r_pass = st.text_input("Bir Şifre Belirleyin:", type="password")
            
            if st.form_submit_button("Kayıt Ol", use_container_width=True):
                # E-posta veritabanında var mı kontrolü
                check_email = supabase.table("kullanicilar").select("*").eq("eposta", r_email.lower()).execute()
                
                # Numara kontrolü (Sadece numara yazılmışsa)
                check_no_exists = False
                if r_no.strip():
                    res_no = supabase.table("kullanicilar").select("*").eq("ogrenci_no", r_no).execute()
                    if len(res_no.data) > 0:
                        check_no_exists = True

                if check_no_exists:
                    st.error("Bu öğrenci numarası zaten sistemde kayıtlı!")
                elif len(check_email.data) > 0:
                    st.error("Bu e-posta zaten kullanımda!")
                elif "@" not in r_email:
                    st.error("Geçerli bir e-posta adresi girin!")
                elif len(r_name) < 2 or len(r_pass) < 3:
                    st.error("Lütfen adınızı ve şifrenizi girin!")
                else:
                    # KAYIT: Boş numara None olarak gider
                    supabase.table("kullanicilar").insert({
                        "ogrenci_no": r_no.strip() if r_no.strip() else None, 
                        "ad_soyad": r_name, 
                        "eposta": r_email.lower(), 
                        "sifre": r_pass
                    }).execute()
                    st.success("Kayıt başarılı! Şimdi e-postanız ile giriş yapabilirsiniz.")
    st.stop()

# --- ANA UYGULAMA ---
ogrenci = st.session_state.user_info

with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/1046/1046269.png", width=150)
    st.title(f"🎓 Merhaba, {ogrenci['ad_soyad'].split()[0]}")
    st.caption(f"Öğrenci No: {ogrenci['ogrenci_no']}")  # Öğrenci numarası geri geldi!
    st.divider()
    
    exp_title = st.selectbox("Konu Başlığı:", [
        "Atom ve Periyodik Sistem", "Kimyasal Bağlar ve Molekül Geometrisi", "Kimyasal Hesaplamalar",
        "Maddenin Halleri", "Sıvı Çözeltiler ve Çözünürlük", "Kimyasal Tepkimelerde Enerji",
        "Kimyasal Tepkimelerde Hız", "Kimyasal Tepkimelerde Denge", "Asitler-Bazlar-Tuzlar", "Elektrokimya"
    ])
    if st.button("✅ Seçimi Onayla ve Başla", use_container_width=True):
        st.session_state.current_subject = exp_title
        st.session_state.messages = [{"role": "assistant", "content": f"👋 Merhaba {ogrenci['ad_soyad']}! **{exp_title}** konusuna hoş geldin! Kafana takılan nedir?"}]
        st.rerun()
        
    st.divider()

    c1, c2 = st.columns(2)
    with c1:
        if st.button("➕ Yeni Sohbet", use_container_width=True):
            st.session_state.messages = [{"role": "assistant", "content": f"👋 Tertemiz yeni bir sayfa açtık! Konumuz: **{st.session_state.current_subject}**."}]
            st.rerun()
    with c2:
        if st.button("💾 Sohbeti Kaydet", use_container_width=True):
            if len(st.session_state.messages) > 0:
                supabase.table("sohbetler").insert({"ogrenci_no": ogrenci['ogrenci_no'], "konu": st.session_state.current_subject, "mesajlar": st.session_state.messages}).execute()
                st.toast("Sohbet başarıyla kaydedildi!", icon="✅")

    st.divider()

    # GEÇMİŞ SOHBETLER GERİ GELDİ!
    with st.expander("📂 Geçmiş Sohbetlerim", expanded=False):
        past_chats_res = supabase.table("sohbetler").select("*").eq("ogrenci_no", ogrenci['ogrenci_no']).order('kayit_tarihi', desc=True).execute()
        past_chats = past_chats_res.data
        if len(past_chats) == 0:
            st.info("Kayıtlı sohbet yok.")
        else:
            for chat in past_chats:
                tarih = chat['kayit_tarihi'].split("T")[0]
                if st.button(f"🕒 {tarih} - {chat['konu']}", key=f"chat_{chat['id']}", use_container_width=True):
                    st.session_state.messages = chat['mesajlar']
                    st.session_state.current_subject = chat['konu']
                    st.rerun()

    st.divider()

    # WORD OLUŞTURMA VE MAİL
    if len(st.session_state.messages) > 0:
        doc = Document()
        doc.add_heading('ChemMind AI - Öğrenci Sohbet Raporu', 0)
        doc.add_paragraph(f"Öğrenci Adı: {ogrenci['ad_soyad']}")
        doc.add_paragraph(f"Öğrenci No: {ogrenci['ogrenci_no']}")
        doc.add_paragraph(f"Çalışılan Konu: {st.session_state.current_subject}")
        doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('Sohbet Geçmişi', level=1)
        for msg in st.session_state.messages:
            p = doc.add_paragraph()
            p.add_run(f"{'Öğrenci' if msg['role'] == 'user' else 'ChemMind AI'}: ").bold = True
            p.add_run(msg["content"])
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button("📄 Word Olarak İndir", data=buffer, file_name=f"{ogrenci['ad_soyad']}_Raporu.docx", use_container_width=True)
        
        if st.button("📧 Hocama Gönder", use_container_width=True):
            buffer.seek(0)
            ok, msg = mail_gonder(buffer, ogrenci['ad_soyad'], st.session_state.current_subject)
            if ok: st.toast(msg, icon="✅")
            else: st.error(msg)

    st.divider()
    
    if st.button("🚪 Çıkış Yap", use_container_width=True):
        st.session_state.force_logout = True
        st.session_state.logged_in = False
        st.session_state.user_info = None
        st.session_state.messages = []
        cookie_manager.delete("chem_user")
        st.rerun()

# --- SİSTEM TALİMATI ---
sistem_promptu = f"Sen destekleyici bir Kimya Öğretmenisin. Öğrenci: {ogrenci['ad_soyad']}. Konu: '{st.session_state.current_subject}'. Cevapları doğrudan vermek yerine sorgulat."

# --- SOHBET EKRANI ---
st.title("🔬 ChemMind AI: İnteraktif Kimya Laboratuvarı")  # Başlık düzeltildi!

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Mesaj kutusu yazısı düzeltildi!
if prompt := st.chat_input("ChemMind' a bir şey sor..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        with st.spinner("ChemMind düşünüyor..."):
            try:
                response = model.generate_content(f"{sistem_promptu}\n\nÖğrenci Mesajı: {prompt}")
                st.markdown(response.text)
                st.session_state.messages.append({"role": "assistant", "content": response.text})
            except Exception as e:
                st.error("Google Gemini geçiolarak meşgul. Lütfen 1-2 dakika sonra tekrar deneyin.")
